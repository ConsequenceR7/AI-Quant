# -*- coding: utf-8 -*-
"""
量化交易工作坊 TASK1 - 报告生成脚本
生成符合格式要求的 Word 文档（可导出为 PDF）

格式要求：
  宋体（SimSun），五号字（10.5pt）
  1.5 倍行距
  0 倍段间距
  文字两端对齐
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

# ============================================================
# 配置区
# ============================================================
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
STOCK_CODE = "002747"
STOCK_NAME = "埃斯顿"
AUTHOR_NAME = "李菓"

FONT_NAME = "SimSun"        # 宋体
FONT_SIZE = Pt(10.5)         # 五号字
LINE_SPACING = 1.5           # 1.5 倍行距


def add_formatted_paragraph(doc, text, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """添加格式化段落"""
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = FONT_SIZE
    run.font.bold = bold
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return p


def add_heading_formatted(doc, text, level=1):
    """添加格式化标题: level=1 一级, level=2 二级, level=3 三级"""
    if level == 1:
        p = add_formatted_paragraph(doc, text, bold=True,
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER)
        p.runs[0].font.size = Pt(14)
    elif level == 2:
        p = add_formatted_paragraph(doc, text, bold=True,
                                    alignment=WD_ALIGN_PARAGRAPH.LEFT)
        p.runs[0].font.size = Pt(12)
    elif level == 3:
        p = add_formatted_paragraph(doc, text, bold=True,
                                    alignment=WD_ALIGN_PARAGRAPH.LEFT)
        p.runs[0].font.size = Pt(10.5)
    return p


def add_empty_line(doc):
    """添加空行"""
    add_formatted_paragraph(doc, "", alignment=WD_ALIGN_PARAGRAPH.CENTER)


def create_report():
    """创建完整任务报告"""
    doc = Document()

    # ---- 设置默认样式 ----
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # ---- 设置页面边距 ----
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # ================================================================
    # 封面 / 标题
    # ================================================================
    add_empty_line(doc)
    add_heading_formatted(doc, "量化交易工作坊 TASK1", level=1)
    add_heading_formatted(doc, "量化交易初体验：从零搭建数据引擎", level=1)
    add_empty_line(doc)
    add_formatted_paragraph(doc, f"姓名：{AUTHOR_NAME}",
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_empty_line(doc)

    # ================================================================
    # 问题一
    # ================================================================
    add_heading_formatted(doc,
        "一、相较于传统手工操作交易的方法，量化交易有哪些优势？", level=2)

    q1 = [
        "量化交易（Quantitative Trading）是指利用数学模型、统计方法和计算机程序，"
        "对金融市场数据进行系统化分析，并据此做出交易决策的一种交易方式。"
        "相较于传统的手工操作交易方法，量化交易具有以下显著优势：",

        "1. 系统性与纪律性。量化交易基于预先设定的规则和模型进行决策，"
        "避免了人类情绪（如恐惧、贪婪、过度自信）对交易判断的干扰。"
        "传统手工交易中，交易者常常因市场短期波动而做出非理性的买卖决策，"
        "而量化系统严格遵循策略逻辑，确保每一笔交易都在既定框架内执行。",

        "2. 高效性与实时性。计算机程序可以在毫秒级别内处理海量市场数据，"
        "同时对数百只甚至数千只股票进行监控和分析。传统手工交易受限于"
        "人类的信息处理速度，难以在短时间内捕捉稍纵即逝的交易机会。"
        "量化系统能够实时扫描全市场，发现符合策略条件的交易信号并迅速执行。",

        "3. 可回测性与可验证性。量化交易策略可以利用历史数据进行回测"
        "（Backtesting），在投入真实资金之前评估策略的历史表现、"
        "风险收益特征和稳定性。传统手工交易方法难以进行严格的量化回测，"
        "策略的有效性往往依赖于交易者的主观判断和经验。",

        "4. 风险管理的精细化。量化交易可以对投资组合进行精确的风险度量"
        "（如 VaR、最大回撤、夏普比率等），并通过程序化方式设置止损、"
        "止盈和仓位管理规则。这种数量化的风险控制手段远优于传统手工"
        "交易中凭经验设定的风控方式。",

        "5. 分散化与规模效应。量化系统可以同时管理多个策略、多个品种和"
        "多个市场，实现真正意义上的分散化投资。传统手工交易者受精力限制，"
        "通常只能关注有限数量的标的。量化交易的规模效应能够通过分散投资"
        "降低非系统性风险，提高整体投资组合的稳健性。",

        "6. 可复制性与可扩展性。一个成熟的量化策略可以很容易地在不同市场、"
        "不同时间框架下进行复制和扩展。量化模型一旦开发完成，其维护和迭代"
        "的成本相对有限，而传统手工交易依赖于个别交易者的经验和技术，"
        "难以进行规模化和标准化的复制。",

        "综上所述，量化交易在系统性、效率、风险管理、可验证性和可扩展性"
        "等方面均明显优于传统手工操作交易方法，这也是其近年来在全球金融"
        "市场中得到广泛应用的根本原因。",
    ]
    for text in q1:
        add_formatted_paragraph(doc, text)

    # ================================================================
    # 问题二：基本概念解释
    # ================================================================
    add_heading_formatted(doc, "二、基本概念解释", level=2)

    # --- K线 ---
    add_heading_formatted(doc, "2.1 K线（K-Line / Candlestick）", level=3)

    kline = [
        "K线，又称蜡烛图（Candlestick Chart）或日本线，起源于18世纪日本的"
        "大米交易市场，由日本商人本间宗久创立，后经史蒂夫.尼森（Steve Nison）"
        "引入西方金融界，现已成为全球金融市场最广泛使用的价格图表形式之一。",

        "一根K线包含四个核心价格信息：开盘价（Open）、收盘价（Close）、"
        "最高价（High）和最低价（Low），因此也被称为OHLC数据。"
        "K线的构成分为实体（Body）和影线（Shadow/Wick）两部分："
        "开盘价与收盘价之间的矩形区域称为「实体」，"
        "如果收盘价高于开盘价，实体通常以红色或空心表示（阳线），代表价格上涨；"
        "如果收盘价低于开盘价，实体通常以绿色或实心表示（阴线），代表价格下跌。"
        "实体上方延伸到最高价的细线称为「上影线」，"
        "下方延伸到最低价的细线称为「下影线」。",

        "K线根据时间周期的不同，可分为日K线、周K线、月K线、"
        "60分钟K线、15分钟K线等多种类型。通过观察K线的形态、组合和趋势，"
        "交易者可以分析市场多空力量的对比，判断价格的短期和中期走势。"
        "常见的K线形态包括十字星（Doji，多空力量均衡）、"
        "锤子线（Hammer，潜在反转信号）、"
        "吞没形态（Engulfing，趋势反转信号）和"
        "三只乌鸦（Three Black Crows，持续下跌信号）等。",

        "K线在量化交易中是技术分析最重要的基础数据之一。"
        "量化策略通常直接使用K线的OHLC数据计算各类技术指标"
        "（如移动平均线、RSI、MACD等），或利用K线形态识别算法自动检测特定的"
        "价格形态，从而实现自动化的交易信号生成。",
    ]
    for text in kline:
        add_formatted_paragraph(doc, text)

    # --- 基本面 ---
    add_heading_formatted(doc, "2.2 基本面（Fundamentals）", level=3)

    fundamental = [
        "基本面分析（Fundamental Analysis）是一种通过评估影响资产内在价值的"
        "经济、财务和行业因素，来判断资产是否被高估或低估的投资分析方法。"
        "基本面分析的核心假设是：市场价格虽然短期内可能偏离内在价值，"
        "但长期来看价格终将回归价值。",

        "在股票投资中，基本面分析通常涵盖三个层面："
        "（1）宏观经济分析——包括GDP增长、通货膨胀率、利率政策、货币政策、"
        "就业数据、国际贸易等宏观经济指标，这些因素决定了整个市场的大环境；"
        "（2）行业分析——研究公司所处行业的发展阶段、竞争格局、市场规模、"
        "监管政策、技术趋势等，判断行业的成长性和盈利前景；"
        "（3）公司分析——深入分析公司的财务报表（资产负债表、利润表、现金流量表），"
        "计算关键财务指标如市盈率（P/E）、市净率（P/B）、净资产收益率（ROE）、"
        "资产负债率、毛利率、净利润增长率等，评估公司的盈利能力、成长性、"
        "偿债能力和经营效率。",

        "基本面分析在量化交易中通常体现为「因子投资」（Factor Investing）——"
        "将基本面的各项指标量化为可计算、可测试的因子"
        "（如价值因子、成长因子、质量因子等），通过统计方法验证因子与股票未来"
        "收益之间的关系，构建基于基本面因子的多因子选股模型。",

        "基本面分析的优势在于其逻辑基础扎实，适合中长期投资决策；"
        "其局限性在于信息获取具有滞后性，且对公司和管理层的定性判断难以完全量化。",
    ]
    for text in fundamental:
        add_formatted_paragraph(doc, text)

    # --- 技术面 ---
    add_heading_formatted(doc, "2.3 技术面（Technicals）", level=3)

    technical = [
        "技术分析（Technical Analysis）是一种通过研究历史市场价格、成交量和其他"
        "交易数据，来预测未来价格走势的分析方法。技术分析建立在三个基本假设之上："
        "（1）市场行为包容消化一切信息（Market Action Discounts Everything）——"
        "所有影响价格的因素（基本面、心理、政治等）都已反映在价格之中；"
        "（2）价格以趋势方式演变（Prices Move in Trends）——价格运动具有一定的"
        "惯性，一旦趋势形成，价格更有可能继续沿趋势方向运行而非逆转；"
        "（3）历史会重演（History Repeats Itself）——由于人类心理和行为模式的"
        "相对稳定性，历史上出现的价格形态和趋势模式会在未来重复出现。",

        "技术分析的主要工具和方法包括："
        "（1）趋势分析——通过趋势线、通道线和移动平均线等工具识别和跟踪价格趋势"
        "的方向和强度；"
        "（2）形态分析——识别图表上的特定价格形态，如头肩顶/底、双顶/底、"
        "三角形整理、旗形等经典形态，判断趋势的延续或反转；"
        "（3）技术指标分析——利用数学公式对价格和成交量进行计算，生成各类技术指标，"
        "如相对强弱指数（RSI）、随机指标（KDJ）、指数平滑异同移动平均线（MACD）、"
        "布林带（Bollinger Bands）、平均趋向指数（ADX）等；"
        "（4）成交量分析——结合成交量变化验证价格趋势的有效性，经典的量价关系包括"
        "「量增价涨」（上升趋势健康）和「量增价跌」（下跌趋势强劲）等。",

        "在量化交易中，技术面是策略开发的核心数据来源之一。"
        "技术指标可以被精确地公式化和程序化，通过回测系统验证其有效性。"
        "许多经典的量化策略（如趋势跟踪策略、均值回归策略、动量策略、突破策略等）"
        "都以技术面分析为主要理论基础。技术分析的优势在于其客观性、可量化性和实时性；"
        "其局限性在于过度依赖历史统计规律，在市场结构发生根本性变化时可能失效。",

        "需要指出的是，基本面分析和技术面分析并非相互排斥，而是相辅相成的关系。"
        "在实际的量化交易实践中，许多成功的策略会同时融合基本面因子和技术面因子，"
        "形成「基本面选股+技术面择时」的复合策略框架。例如，利用基本面因子筛选出"
        "优质股票池，再利用技术指标确定最佳的买入和卖出时机。",
    ]
    for text in technical:
        add_formatted_paragraph(doc, text)

    # ================================================================
    # 问题三：数据引擎实践
    # ================================================================
    add_heading_formatted(doc, "三、量化数据引擎实践：股票数据获取与可视化", level=2)

    add_formatted_paragraph(doc,
        f"本部分基于 Python 金融数据接口，编程实现了对{STOCK_NAME}"
        f"（{STOCK_CODE}）过去一年内每个交易日交易数据的获取、可视化和存储。"
        "数据采集使用 Baostock（http://baostock.com）和 Tushare"
        "（https://www.tushare.pro/）两大国内广受认可的免费金融数据平台，"
        "为量化交易爱好者和研究者提供了丰富、高质量的A股数据支持。"
    )

    add_heading_formatted(doc, "3.1 数据获取方法", level=3)

    method = [
        "首先，在 Tushare 官网（https://www.tushare.pro/）注册账号并获取个人"
        "API token，完成数据平台的接入准备。同时，为了确保数据获取的稳定性，"
        "本项目也集成了 Baostock 证券数据平台（无需注册，直接通过 Python SDK 调用），"
        "通过 query_history_k_data_plus() 接口可以高效地获取各类金融数据。",
        "接口可以高效地获取各类金融数据。",

        f"本次任务使用 query_history_k_data_plus() 接口获取{STOCK_NAME}（{STOCK_CODE}）"
        "的日线行情数据（前复权），"
        "数据字段包括：交易日期（trade_date）、开盘价（open）、最高价（high）、"
        "最低价（low）、收盘价（close）、前收盘价（pre_close）、涨跌额（change）、"
        "涨跌幅（pct_chg）、成交量（vol）和成交额（amount）。数据时间范围为过去一年"
        "（约 365 个自然日，覆盖约 240-250 个交易日）。",

        "获取数据后，首先按交易日期进行升序排列，将日期字符串转换为 datetime 格式"
        "以便后续时间序列处理和可视化。同时生成数据摘要报告，包括交易日数量、"
        "价格范围、日均成交量和成交额、区间涨跌幅等关键统计信息。",
    ]
    for text in method:
        add_formatted_paragraph(doc, text)

    add_heading_formatted(doc, "3.2 收盘价曲线图分析", level=3)

    add_formatted_paragraph(doc,
        "根据获取的日线数据，绘制了包含收盘价走势和成交量分布的综合图表（见下图）。"
        "图表分为上下两个面板：上方面板展示每日收盘价走势及 20 日/60 日移动均线，"
        "下方面板展示每日成交量分布（红色柱表示当日上涨、蓝色柱表示当日下跌）。"
    )

    # 插入图表
    chart_path = os.path.join(
        OUTPUT_DIR, f"{STOCK_CODE.replace('.', '_')}_close_price.png")
    if os.path.exists(chart_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(chart_path, width=Inches(5.5))
        add_formatted_paragraph(doc,
            f"图1：{STOCK_NAME}({STOCK_CODE}) 每日收盘价走势及成交量分布图",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
    else:
        add_formatted_paragraph(doc,
            f"（注：请先运行 stock_data_engine.py 脚本生成图表文件，"
            f"图表将自动插入此处。预期图表文件：{chart_path}）"
        )

    chart_analysis = [
        "从上方的收盘价走势图中可以观察到以下几点：",

        "（1）整体趋势：通过收盘价曲线和移动平均线（20日均线和60日均线）的相对位置，"
        "可以判断该股票在过去一年的整体走势方向。当短期均线（20日）位于长期均线（60日）"
        "上方时，通常意味着股价处于上升趋势中（即「金叉」信号）；反之则可能处于下跌趋势"
        "（即「死叉」信号）。图中标注了区间内的最高收盘价和最低收盘价位置，"
        "便于直观了解价格波动范围。",

        "（2）价格波动性：收盘价曲线的波动幅度反映了该股票的价格波动率特征。"
        "曲线较为平缓的区间对应低波动期（市场处于盘整阶段），而曲线陡峭上升或下降的"
        "区间对应高波动期（市场出现明确的趋势性行情）。",

        "（3）成交量分析：下方成交量柱状图展示了每日的成交量变化。"
        "通常情况下，价格上涨伴随成交量放大（「量价齐升」）表明上涨趋势较为健康，"
        "得到了市场资金的认可；而价格下跌伴随成交量放大则可能预示着恐慌性抛售或趋势加速。"
        "成交量萎缩期间通常对应市场的观望情绪或盘整阶段。",

        "通过以上图表分析，可以对该股票过去一年的市场表现形成直观、全面的认识，"
        "为后续的量化策略开发和回测提供重要参考。",
    ]
    for text in chart_analysis:
        add_formatted_paragraph(doc, text)

    add_heading_formatted(doc, "3.3 数据存储", level=3)

    csv_text = [
        "最后，将获取的全部日线交易数据保存为 CSV（Comma-Separated Values）格式文件。"
        "CSV 是一种通用、轻量级的数据交换格式，可以被 Excel、Python（pandas）、R、"
        "MATLAB 等各类数据分析工具直接读取，便于在后续任务中进行策略回测、"
        "数据清洗和特征工程等工作。",

        "保存的 CSV 文件使用 UTF-8 with BOM 编码，确保在 Microsoft Excel 中直接打开时"
        "中文列名能够正确显示。数据字段包括：交易日期、开盘价、最高价、最低价、收盘价、"
        "前收盘价、涨跌额、涨跌幅(%)、成交量(手)和成交额(千元)，共计 10 个字段，"
        "覆盖了日线行情分析的完整信息维度。",
    ]
    for text in csv_text:
        add_formatted_paragraph(doc, text)

    # ================================================================
    # 总结
    # ================================================================
    add_heading_formatted(doc, "四、总结", level=2)

    summary = [
        "通过本次 TASK1 任务的实践，我对量化交易的基本概念和数据结构有了初步的"
        "认知和理解。量化交易凭借其系统性、高效性和可验证性等优势，正在重塑现代"
        "金融市场的基本运作方式。K线、基本面和技术面作为量化交易的三大核心概念，"
        "分别从价格行为、内在价值和统计规律三个维度提供了分析金融市场的不同视角。",

        "在实践环节中，我成功搭建了第一个量化数据引擎：通过 Python 开源数据接口获取了"
        "沪深A股的真实日线交易数据，使用 Python 的 matplotlib 库绘制了收盘价走势图和"
        "成交量分布图，并将数据以 CSV 格式保存以便后续任务使用。"
        "这个过程让我深刻体会到：可靠的数据获取和高效的数据处理是量化交易的基石，"
        "只有建立了稳定、规范的数据管道，后续的策略开发和回测工作才能有据可依。",

        "在接下来的任务中，我期待基于本次搭建的数据引擎，进一步学习量化策略的开发、"
        "回测和优化方法，逐步构建起完整的量化交易知识体系。",
    ]
    for text in summary:
        add_formatted_paragraph(doc, text)

    # ================================================================
    # 保存文档
    # ================================================================
    docx_path = os.path.join(OUTPUT_DIR, f"{AUTHOR_NAME}+TASK1.docx")
    doc.save(docx_path)
    print(f"OK 报告文档已保存至: {docx_path}")
    print()
    print("后续操作：")
    print("  1. 请将文档中的「姓名」替换为你的真实姓名")
    print("  2. 用 Microsoft Word 打开文档，检查格式")
    print("  3. 在 Word 中选择「文件 -> 另存为 -> PDF」导出为 PDF 文件")
    print(f"  4. 确认 PDF 文件名为：{AUTHOR_NAME}+TASK1.pdf")

    return docx_path


if __name__ == "__main__":
    create_report()
